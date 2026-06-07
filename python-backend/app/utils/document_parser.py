"""知识库文档文本提取"""

from io import BytesIO
from pathlib import Path

from docx import Document

ALLOWED_KNOWLEDGE_EXTENSIONS = {".txt", ".md", ".docx"}


def normalize_extension(file_name: str) -> str:
    return Path(file_name).suffix.lower()


def is_allowed_knowledge_file(file_name: str) -> bool:
    ext = normalize_extension(file_name)
    if ext == ".doc":
        return False
    return ext in ALLOWED_KNOWLEDGE_EXTENSIONS


def extension_to_file_type(file_name: str) -> str:
    ext = normalize_extension(file_name)
    if ext.startswith("."):
        return ext[1:]
    return ext


def extract_text_from_bytes(file_name: str, file_bytes: bytes) -> str:
    """从上传文件中提取纯文本，供分块与向量化。"""
    ext = normalize_extension(file_name)
    if ext == ".doc":
        raise ValueError("不支持旧版 .doc，请在 Word 中另存为 .docx 后上传")
    if ext not in ALLOWED_KNOWLEDGE_EXTENSIONS:
        raise ValueError("仅支持 .txt / .md / .docx 文件")

    if ext in {".txt", ".md"}:
        text = file_bytes.decode("utf-8", errors="ignore").strip()
    elif ext == ".docx":
        text = _extract_docx_text(file_bytes)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")

    if not text:
        raise ValueError("文件内容为空或无法解析出文本")
    return text


def _extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if line:
            parts.append(line)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts).strip()
